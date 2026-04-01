from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from reports.base_generator import BaseReportGenerator
from reports.models import ReportData

_HEADER_COLUMNS = [
    "modelo", "marca", "status task", "detalle",
    "scheduled_at", "started_at", "end_at",
]

_DATE_FMT = "%Y-%m-%d %H:%M"


def _fmt(value: datetime | None) -> str:
    return value.strftime(_DATE_FMT) if value else "-"


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        from lxml import etree
        shd = etree.SubElement(tcPr, qn("w:shd"))
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


class WordReportGenerator(BaseReportGenerator):

    def generate(self, data: ReportData) -> bytes:
        doc = Document()

        # ── Title ──
        title = doc.add_heading("MASSIVE TASK PROGRAM", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # ── Metadata ──
        meta_items = [
            ("nombre tarea", data.task_name),
            ("total devices", str(data.total_devices)),
            ("processed devices", str(data.processed_devices)),
            ("failed devices", str(data.failed_devices)),
            ("group name", data.group_name),
        ]
        for label, value in meta_items:
            para = doc.add_paragraph()
            run_label = para.add_run(f"{label}: ")
            run_label.bold = True
            para.add_run(value)

        doc.add_paragraph()

        # ── Table ──
        num_cols = 1 + len(_HEADER_COLUMNS)  # serial_number col + data cols
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = "Table Grid"

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ""
        for i, col_name in enumerate(_HEADER_COLUMNS, start=1):
            hdr_cells[i].text = col_name
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(hdr_cells[i], "2C3E50")

        _set_cell_bg(hdr_cells[0], "2C3E50")

        # Data rows
        for row in data.rows:
            cells = table.add_row().cells
            cells[0].text = row.serial_number
            cells[1].text = row.model or "-"
            cells[2].text = row.manufacturer or "-"
            cells[3].text = row.task_status
            cells[4].text = row.detail or "-"
            cells[5].text = _fmt(row.scheduled_at)
            cells[6].text = _fmt(row.started_at)
            cells[7].text = _fmt(row.end_at)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
